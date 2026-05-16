import io
import os
import re
import time
import asyncio
import logging
import datetime
from typing import Optional

import httpx
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)

app = FastAPI(title="ISDP Platform")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

IBM_API_KEY    = os.getenv("IBM_API_KEY")
WXO_URL        = os.getenv("WXO_URL", "https://api.eu-de.watson-orchestrate.cloud.ibm.com")
ENVIRONMENT_ID = os.getenv("ENVIRONMENT_ID", "draft")

# ── IAM token cache ───────────────────────────────────────────────────────────
_token_cache: dict = {"token": None, "expires_at": 0.0}


async def get_token() -> str:
    if _token_cache["token"] and time.time() < _token_cache["expires_at"] - 60:
        return _token_cache["token"]
    if not IBM_API_KEY:
        raise HTTPException(500, "IBM_API_KEY not set")
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.post(
            "https://iam.cloud.ibm.com/identity/token",
            data={"grant_type": "urn:ibm:params:oauth:grant-type:apikey", "apikey": IBM_API_KEY},
            headers={"Content-Type": "application/x-www-form-urlencoded"},
        )
        if not r.is_success:
            raise HTTPException(401, f"IAM token exchange failed: {r.text[:300]}")
        d = r.json()
        _token_cache["token"] = d["access_token"]
        _token_cache["expires_at"] = time.time() + d["expires_in"]
        logger.info("IAM token refreshed (expires in %ss)", d["expires_in"])
        return _token_cache["token"]


# ── Models ────────────────────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    message: str
    agent_id: str
    thread_id: Optional[str] = None


class DocxRequest(BaseModel):
    project_name: str
    agent_outputs: list[dict]


# ── File upload & text extraction ─────────────────────────────────────────────

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    content = await file.read()
    filename = file.filename or "document"
    text = ""

    try:
        if filename.lower().endswith(".pdf"):
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            text = "\n\n".join(
                page.extract_text() for page in reader.pages if page.extract_text()
            )
        elif filename.lower().endswith((".docx", ".doc")):
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif filename.lower().endswith((".txt", ".md")):
            text = content.decode("utf-8", errors="replace")
        else:
            text = content.decode("utf-8", errors="replace")
    except Exception as e:
        logger.warning("Could not extract text from %s: %s", filename, e)
        text = f"[Could not extract text: {e}]"

    if len(text) > 8000:
        text = text[:8000] + "\n\n[... document truncated ...]"

    logger.info("Uploaded %s (%d chars extracted)", filename, len(text))
    return {"filename": filename, "text": text, "chars": len(text)}


# ── WatsonX Orchestrate run / poll ────────────────────────────────────────────

async def _start_run(
    client: httpx.AsyncClient,
    headers: dict,
    agent_id: str,
    message: str,
    thread_id: Optional[str],
    environment_id: Optional[str],
) -> dict:
    body: dict = {
        "agent_id": agent_id,
        "message": {"role": "user", "content": message},
    }
    if environment_id:
        body["environment_id"] = environment_id
    if thread_id:
        body["thread_id"] = thread_id
    r = await client.post(f"{WXO_URL}/v1/orchestrate/runs", headers=headers, json=body)
    return r


async def run_wxo(
    client: httpx.AsyncClient,
    headers: dict,
    agent_id: str,
    message: str,
    thread_id: Optional[str],
) -> tuple[str, str]:
    # Try with configured environment_id first; fall back to no environment_id
    # if WXO says the agent is not found in that environment.
    r = await _start_run(client, headers, agent_id, message, thread_id, ENVIRONMENT_ID)

    if not r.is_success:
        err_text = r.text
        logger.warning("Run with env=%s failed %s: %s", ENVIRONMENT_ID, r.status_code, err_text[:300])
        # If the error mentions environment, retry without environment_id
        if "environment" in err_text.lower() or r.status_code == 500:
            logger.info("Retrying without environment_id for agent %s", agent_id)
            r = await _start_run(client, headers, agent_id, message, thread_id, None)
        if not r.is_success:
            logger.error("Run failed %s: %s", r.status_code, r.text[:800])
            raise HTTPException(r.status_code, f"Run start failed: {r.text[:800]}")

    run_data = r.json()
    current_run_id: str = run_data["run_id"]
    returned_thread_id: str = run_data["thread_id"]
    accumulated_text = ""

    logger.info("Run started: %s  thread: %s", current_run_id, returned_thread_id)

    _FLOW_MARKER = "A new flow has started"

    for i in range(400):
        await asyncio.sleep(1)
        poll = await client.get(
            f"{WXO_URL}/v1/orchestrate/runs/{current_run_id}", headers=headers
        )
        if not poll.is_success:
            raise HTTPException(poll.status_code, f"Poll failed: {poll.text[:200]}")

        result = poll.json()
        status: str = result["status"]
        logger.info("Poll %d (run %s) status: %s", i, current_run_id, status)

        if status == "failed":
            raise HTTPException(500, f"Run failed: {result.get('last_error', 'unknown')}")

        if status == "completed":
            content = result["result"]["data"]["message"]["content"]
            chunk = " ".join(c.get("text", "") for c in content if isinstance(c, dict))
            next_run_id: Optional[str] = result["result"].get("next_run_id")
            logger.info("Completed chunk (len=%d) next_run_id=%s", len(chunk), next_run_id)

            # WXO skill-flow agents emit a transient "A new flow has started" message
            # when the flow is still running. Only accumulate real content.
            if _FLOW_MARKER not in chunk:
                accumulated_text += chunk

            if next_run_id:
                current_run_id = next_run_id
                logger.info("Continuing with next_run_id: %s", next_run_id)
            else:
                if not accumulated_text and _FLOW_MARKER in chunk:
                    # Flow agent returned only the marker with no follow-up run — surface it
                    # so the frontend can show a meaningful info message.
                    return returned_thread_id, chunk
                return returned_thread_id, accumulated_text

    raise HTTPException(504, "Run timed out after 400 seconds")


# ── API endpoints ─────────────────────────────────────────────────────────────

@app.post("/api/chat")
async def chat(req: ChatRequest):
    token = await get_token()
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    async with httpx.AsyncClient(timeout=httpx.Timeout(420.0, connect=10.0)) as client:
        thread_id, response_text = await run_wxo(
            client, headers, req.agent_id, req.message, req.thread_id
        )
    return {"response": response_text, "thread_id": thread_id, "agent_id": req.agent_id}


@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "wxo_url": WXO_URL,
        "environment_id": ENVIRONMENT_ID,
        "api_key_set": bool(IBM_API_KEY),
    }


@app.get("/api/debug/agents")
async def debug_agents():
    """List all agents available in the WXO instance so we can verify agent IDs."""
    token = await get_token()
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.get(f"{WXO_URL}/v1/agents", headers=headers)
        if r.is_success:
            return {"source": "/v1/agents", "data": r.json()}
        r2 = await client.get(f"{WXO_URL}/v1/orchestrate/agents", headers=headers)
        if r2.is_success:
            return {"source": "/v1/orchestrate/agents", "data": r2.json()}
        return {
            "error": f"Could not list agents.",
            "body1": r.text[:500],
            "body2": r2.text[:500],
        }


# ── DOCX generation helpers ───────────────────────────────────────────────────

GLOSSARY = {
    "ISDP": "Informationssicherheits- und Datenschutzkonzept",
    "IAM": "Identity and Access Management (Identitäts- und Zugriffsverwaltung)",
    "MFA": "Multi-Faktor-Authentifizierung",
    "SAML": "Security Assertion Markup Language",
    "DPIA": "Data Protection Impact Assessment (Datenschutz-Folgenabschätzung)",
    "DSFA": "Datenschutz-Folgenabschätzung",
    "AVV": "Auftrags-Verarbeitungs-Vertrag",
    "BCM": "Business Continuity Management",
    "SIEM": "Security Information and Event Management",
    "SLA": "Service Level Agreement",
    "RTO": "Recovery Time Objective (Maximale Wiederherstellungszeit)",
    "RPO": "Recovery Point Objective (Maximaler Datenverlust-Zeitraum)",
    "nDSG": "Neues Datenschutzgesetz (Schweiz, in Kraft seit 01.09.2023)",
    "AHV": "Alters- und Hinterlassenenversicherung (CH-Sozialversicherungsnummer)",
    "BSI": "Bundesamt für Sicherheit in der Informationstechnik",
    "ISO": "International Organization for Standardization",
    "NIST": "National Institute of Standards and Technology",
    "CSF": "Cybersecurity Framework (NIST)",
    "TLS": "Transport Layer Security",
    "AES": "Advanced Encryption Standard",
    "RSA": "Rivest–Shamir–Adleman (asymmetrisches Verschlüsselungsverfahren)",
    "VPN": "Virtual Private Network",
    "API": "Application Programming Interface",
    "SaaS": "Software as a Service",
    "SOC": "Security Operations Center",
    "GAP": "Identifizierte Lücke im Sicherheits- oder Compliance-Bereich",
    "RACI": "Responsible, Accountable, Consulted, Informed (Verantwortlichkeitsmatrix)",
}

FB_LABELS = {
    "FB-001": "IAM-Rollen-Matrix & Admin-Workflow",
    "FB-002": "Logging / SIEM – Log-Quellen & Aufbewahrung",
    "FB-003": "Backup & BCM – Plan, RPO/RTO, Test-Restore",
    "FB-004": "Datenschutz & DPIA – AVV, Lösch-Procedures",
    "FB-005": "Sicherheits-Governance – Policy-Referenzen",
    "FB-006": "Risikomanagement – Risiko-Register (ISO 27005)",
    "FB-007": "Architektur – Betriebs- & Patch-Management",
    "FB-008": "IAM Scope-Anpassung (Need-to-Know für Personendaten)",
    "FB-009": "Log-Retention 90 Tage – bestätigt",
    "FB-010": "Backup RPO 4h / RTO 8h – bestätigt",
    "FB-011": "DPIA-Entwurf akzeptiert, AVV noch offen",
}


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_module_break())


def docx_module_break():
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    br = _OE("w:br")
    br.set(_qn("w:type"), "page")
    return br


def _cover_page(doc, project_name: str):
    # Red top bar
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, "C8102E")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("ABRAXAS INFORMATIK AG  ·  ISDP-Plattform")
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Informationssicherheits- und\nDatenschutzkonzept (ISDP)")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    doc.add_paragraph()

    # Project name
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Projekt / System: {project_name}")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x5C, 0x5C, 0x5C)

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata table
    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    meta_data = [
        ("Version", "1.0 – Entwurf"),
        ("Datum", datetime.date.today().strftime("%d.%m.%Y")),
        ("Status", "Entwurf / Draft"),
        ("Klassifikation", "Vertraulich"),
        ("Erstellt durch", "ISDP-Plattform (AI-unterstützt)"),
    ]
    for i, (label, value) in enumerate(meta_data):
        lc = meta.cell(i, 0)
        vc = meta.cell(i, 1)
        _set_cell_bg(lc, "F4F3F0")
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.bold = True
        lr.font.size = Pt(10)
        vp = vc.paragraphs[0]
        vp.add_run(value).font.size = Pt(10)

    doc.add_paragraph()

    # Disclaimer
    d = doc.add_paragraph(
        "Dieses Dokument wurde mit Unterstützung der ISDP-Plattform (KI-gestützt) erstellt "
        "und muss vor der finalen Freigabe durch den zuständigen Informationssicherheitsbeauftragten "
        "und Datenschutzbeauftragten geprüft und unterzeichnet werden."
    )
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.runs[0].font.size = Pt(9)
    d.runs[0].font.color.rgb = RGBColor(0x9A, 0x9A, 0x9A)

    # Page break
    doc.add_page_break()


def _revision_table(doc):
    doc.add_heading("Dokumenten-Kontrolle", level=1)
    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    headers = ["Version", "Datum", "Autor", "Änderung", "Status"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, "C8102E")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
    row1 = ["1.0", datetime.date.today().strftime("%d.%m.%Y"), "ISDP-Plattform", "Erstversion (AI-Draft)", "Entwurf"]
    for i, v in enumerate(row1):
        tbl.cell(1, i).paragraphs[0].add_run(v).font.size = Pt(9)
    doc.add_paragraph()


def _parse_markdown_section(doc, text: str, base_level: int = 1):
    """Convert a markdown string into Word document elements."""
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Tables
        if line.strip().startswith("|") and "|" in line:
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            _render_md_table(doc, tbl_lines)
            continue

        # ATX headings
        m3 = re.match(r"^### (.+)", line)
        m2 = re.match(r"^## (.+)", line)
        m1 = re.match(r"^# (.+)", line)
        mb = re.match(r"^\*\*(.+?)\*\*\s*$", line.strip())
        if m3:
            doc.add_heading(_strip_md(m3.group(1)), level=min(base_level + 2, 4))
            i += 1
            continue
        if m2:
            doc.add_heading(_strip_md(m2.group(1)), level=min(base_level + 1, 3))
            i += 1
            continue
        if m1:
            doc.add_heading(_strip_md(m1.group(1)), level=base_level)
            i += 1
            continue
        if mb:
            doc.add_heading(_strip_md(mb.group(1)), level=min(base_level + 1, 3))
            i += 1
            continue

        # Bullet list
        if re.match(r"^\s*[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, re.sub(r"^\s*[-*] ", "", line))
            i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\s*\d+\. ", "", line))
            i += 1
            continue

        # Horizontal rule → skip
        if re.match(r"^[-–—]{3,}$", line.strip()):
            doc.add_paragraph()
            i += 1
            continue

        # Empty
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_inline(p, line)
        i += 1


def _strip_md(text: str) -> str:
    """Remove markdown formatting characters from a string."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def _add_inline(paragraph, text: str):
    """Add inline-formatted text to a Word paragraph (bold, italic)."""
    parts = re.split(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def _render_md_table(doc, table_lines: list[str]):
    """Render a markdown table as a Word table."""
    data_rows = [l for l in table_lines if not re.match(r"^[\|\s\-:]+$", l.strip())]
    if not data_rows:
        return
    cols = max(len(r.split("|")) - 2 for r in data_rows)
    if cols <= 0:
        return
    tbl = doc.add_table(rows=len(data_rows), cols=cols)
    tbl.style = "Table Grid"
    for row_idx, row_line in enumerate(data_rows):
        cells_text = [c.strip() for c in row_line.split("|")[1:-1]]
        is_header = row_idx == 0
        for col_idx in range(cols):
            cell = tbl.cell(row_idx, col_idx)
            text = cells_text[col_idx] if col_idx < len(cells_text) else ""
            if is_header:
                _set_cell_bg(cell, "E8E6E2")
                p = cell.paragraphs[0]
                r = p.add_run(_strip_md(text))
                r.bold = True
                r.font.size = Pt(9)
            else:
                p = cell.paragraphs[0]
                _add_inline(p, _strip_md(text))
                p.runs[0].font.size = Pt(9) if p.runs else None
    doc.add_paragraph()


def _glossary_section(doc):
    doc.add_heading("Anhang A – Glossar & Abkürzungsverzeichnis", level=1)
    tbl = doc.add_table(rows=len(GLOSSARY) + 1, cols=2)
    tbl.style = "Table Grid"
    # Header
    _set_cell_bg(tbl.cell(0, 0), "C8102E")
    _set_cell_bg(tbl.cell(0, 1), "C8102E")
    for col, label in enumerate(["Abkürzung", "Bedeutung"]):
        p = tbl.cell(0, col).paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
    for row_idx, (abbr, meaning) in enumerate(sorted(GLOSSARY.items()), start=1):
        tbl.cell(row_idx, 0).paragraphs[0].add_run(abbr).font.size = Pt(9)
        tbl.cell(row_idx, 1).paragraphs[0].add_run(meaning).font.size = Pt(9)
    doc.add_paragraph()


def _feedback_register(doc, feedback_items: dict):
    doc.add_heading("Anhang B – Offene Punkte & Feedback-Register", level=1)
    p = doc.add_paragraph(
        "Die folgende Tabelle enthält alle offenen Feedback-Punkte, die vor der finalen Freigabe "
        "des ISDP durch das Projektteam bearbeitet werden müssen."
    )
    p.runs[0].font.size = Pt(10)
    tbl = doc.add_table(rows=len(FB_LABELS) + 1, cols=3)
    tbl.style = "Table Grid"
    headers = ["ID", "Thema", "Status"]
    for col, h in enumerate(headers):
        _set_cell_bg(tbl.cell(0, col), "C8102E")
        r = tbl.cell(0, col).paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
    for row_idx, (fb_id, label) in enumerate(FB_LABELS.items(), start=1):
        tbl.cell(row_idx, 0).paragraphs[0].add_run(fb_id).font.size = Pt(9)
        tbl.cell(row_idx, 1).paragraphs[0].add_run(label).font.size = Pt(9)
        status = feedback_items.get(fb_id, "offen")
        tbl.cell(row_idx, 2).paragraphs[0].add_run(status).font.size = Pt(9)
    doc.add_paragraph()


def build_isdp_docx(project_name: str, agent_outputs: list[dict]) -> io.BytesIO:
    """Build a formatted ISDP Word document from agent outputs."""
    doc = DocxDocument()

    # Page margins (A4)
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)

    # Default body font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # Collect outputs; skip empty / flow-marker placeholders
    _FLOW_MARKER = "A new flow has started"
    output_map: dict[str, str] = {}
    for o in agent_outputs:
        name = o.get("name", "")
        text = o.get("output", "").strip()
        if text and _FLOW_MARKER not in text and len(text) > 30:
            output_map[name] = text

    # 1. Cover page
    _cover_page(doc, project_name)

    # 2. Document control / revision
    _revision_table(doc)
    doc.add_page_break()

    # 3. Workflow status overview
    all_agents = [
        "Document Classifier", "Document Readiness", "Scope Clarification",
        "Reference & Reuse", "ISDP Template", "Standards Relevance",
        "Security Extraction", "Schutzedarfsanalyse", "ISDP Section Mapping",
        "Gap & Conflict Detection", "Clarification Questions", "Standards Coverage",
        "Evidence Validation", "German Draft", "Quality Assurance",
        "Feedback Integration", "Quality Scoring", "Final Review",
    ]
    doc.add_heading("Workflow-Status", level=1)
    status_tbl = doc.add_table(rows=len(all_agents) + 1, cols=2)
    status_tbl.style = "Table Grid"
    for col, lbl in enumerate(["Workflow-Schritt", "Status"]):
        _set_cell_bg(status_tbl.cell(0, col), "C8102E")
        r = status_tbl.cell(0, col).paragraphs[0].add_run(lbl)
        r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(9)
    for row_idx, agent in enumerate(all_agents, start=1):
        status_tbl.cell(row_idx, 0).paragraphs[0].add_run(agent).font.size = Pt(9)
        done = agent in output_map
        status_cell = status_tbl.cell(row_idx, 1)
        sr = status_cell.paragraphs[0].add_run("✓ Abgeschlossen" if done else "– Ausstehend")
        sr.font.size = Pt(9)
        sr.font.color.rgb = RGBColor(0x16, 0x65, 0x34) if done else RGBColor(0x9A, 0x9A, 0x9A)
    doc.add_paragraph()
    doc.add_page_break()

    # Helper: write a section only if output exists
    def _section(heading: str, key: str, level: int = 1, base: int = 2):
        text = output_map.get(key, "")
        if not text:
            return
        doc.add_heading(heading, level=level)
        _parse_markdown_section(doc, text, base_level=base)
        doc.add_page_break()

    # 4. Scope & System Overview
    _section("Scope & System-Übersicht", "Scope Clarification", level=1, base=2)

    # 5. ISDP – Kapitelentwürfe (main body — from German Draft)
    german_draft = output_map.get("German Draft", "")
    doc.add_heading("ISDP – Kapitelentwürfe", level=1)
    if german_draft:
        _parse_markdown_section(doc, german_draft, base_level=2)
    else:
        p = doc.add_paragraph(
            "⚠  Die deutschen Kapitelentwürfe (Schritt «German Draft») stehen noch nicht zur Verfügung. "
            "Dieser Agent ist in WatsonX Orchestrate als Skill-Flow konfiguriert und muss auf "
            "«Conversational» umgestellt werden. Alle anderen verfügbaren Analyseergebnisse sind in "
            "den nachfolgenden Abschnitten dieses Dokuments enthalten."
        )
        p.runs[0].font.color.rgb = RGBColor(0x78, 0x35, 0x0F)
        p.runs[0].font.size = Pt(10)
    doc.add_page_break()

    # 6. ISDP Template / Chapter Structure
    _section("ISDP-Zielstruktur", "ISDP Template", level=1, base=2)

    # 7. Standards & Controls
    _section("Standards-Relevanz", "Standards Relevance", level=1, base=2)
    _section("Standards-Coverage-Check", "Standards Coverage", level=1, base=2)

    # 8. Security Findings
    _section("Security-Findings (extrahiert)", "Security Extraction", level=1, base=2)

    # 9. Schutzbedarfsanalyse
    _section("Schutzbedarfsanalyse", "Schutzedarfsanalyse", level=1, base=2)

    # 10. ISDP Section Mapping
    _section("ISDP-Section-Mapping", "ISDP Section Mapping", level=1, base=2)

    # 11. GAP- and Conflict Register
    _section("GAP- und Konflikt-Register", "Gap & Conflict Detection", level=1, base=2)

    # 12. Clarification Questions
    _section("Klärungsfragen an das Projektteam", "Clarification Questions", level=1, base=2)

    # 13. Evidence Validation
    _section("Evidenz-Validierung", "Evidence Validation", level=1, base=2)

    # 14. Quality Assurance
    _section("Qualitätssicherung (QA-Review)", "Quality Assurance", level=1, base=2)

    # 15. Feedback Integration
    _section("Feedback-Integration", "Feedback Integration", level=1, base=2)

    # 16. Quality Scoring
    _section("Qualitäts-Scoring", "Quality Scoring", level=1, base=2)

    # 17. Final Review
    _section("Final-Review-Gate", "Final Review", level=1, base=2)

    # 18. Reference & Reuse (lower priority — put at end)
    _section("Referenz- und Wiederverwendungsanalyse", "Reference & Reuse", level=1, base=2)

    # 19. Glossary
    _glossary_section(doc)

    # 20. Feedback register (open items)
    _feedback_register(doc, {})

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── DOCX download endpoint ─────────────────────────────────────────────────────

@app.post("/api/generate-docx")
async def generate_docx(data: DocxRequest):
    try:
        buf = build_isdp_docx(data.project_name, data.agent_outputs)
    except Exception as e:
        logger.exception("DOCX generation failed")
        raise HTTPException(500, f"DOCX generation failed: {e}")
    safe_name = re.sub(r"[^\w\-]", "_", data.project_name)
    filename = f"ISDP_{safe_name}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Serve frontend (must be last) ─────────────────────────────────────────────
app.mount("/", StaticFiles(directory="public", html=True), name="static")